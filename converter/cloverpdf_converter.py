#!/usr/bin/env python3
"""Fixed JSONL wrapper around pdf2docx for CloverPDF."""

from __future__ import annotations

import json
import os
import sys
import traceback
from dataclasses import dataclass
from typing import Any


@dataclass(frozen=True)
class Request:
    input: str
    output: str
    password: str | None
    pages: tuple[int, ...]
    ocr_pages: tuple[dict[str, Any], ...] | None


class InputNotFoundError(FileNotFoundError):
    pass


def emit(event_type: str, **values: Any) -> None:
    payload = {"type": event_type, **values}
    sys.stdout.write(json.dumps(payload, ensure_ascii=True) + "\n")
    sys.stdout.flush()


def read_request() -> Request:
    line = sys.stdin.readline()
    if not line:
        raise ValueError("missing_request")
    payload = json.loads(line)
    return Request(
        input=str(payload["input"]),
        output=str(payload["output"]),
        password=payload.get("password"),
        pages=tuple(int(page) for page in payload.get("pages", [])),
        ocr_pages=(
            tuple(payload["ocrPages"])
            if payload.get("ocrPages") is not None
            else None
        ),
    )


def error_code(error: Exception) -> str:
    message = str(error).lower()
    if "require password" in message:
        return "password_required"
    if "incorrect password" in message:
        return "incorrect_password"
    if "no parsed pages" in message or "no_pages" in message:
        return "no_pages"
    if isinstance(error, InputNotFoundError):
        return "input_not_found"
    return "conversion_failed"


def convert(request: Request) -> None:
    if not request.pages or any(page < 0 for page in request.pages):
        raise ValueError("no_pages")
    os.makedirs(os.path.dirname(request.output), exist_ok=True)
    emit("started")
    if request.ocr_pages is not None:
        convert_ocr(request)
        return
    if not os.path.isfile(request.input):
        raise InputNotFoundError(request.input)
    from pdf2docx import Converter

    emit("progress", progress=0.05, phase="opening")
    converter = Converter(request.input, request.password or "")
    try:
        kwargs: dict[str, Any] = {"multi_processing": False}
        emit("progress", progress=0.15, phase="parsing")
        converter.convert(request.output, pages=list(request.pages), **kwargs)
    finally:
        converter.close()
    if not os.path.isfile(request.output) or os.path.getsize(request.output) == 0:
        raise RuntimeError("missing_output")
    emit("progress", progress=1.0, phase="completed")
    emit("completed", output=request.output)


def convert_ocr(request: Request) -> None:
    pages = request.ocr_pages or ()
    build_ocr_document(pages, request.output)
    for page_offset, _ in enumerate(pages):
        emit(
            "progress",
            progress=0.7 + ((page_offset + 1) / max(len(pages), 1)) * 0.3,
            phase="writing_ocr",
        )
    if not os.path.isfile(request.output) or os.path.getsize(request.output) == 0:
        raise RuntimeError("missing_output")
    emit("completed", output=request.output)


def build_ocr_document(pages: tuple[dict[str, Any], ...], output: str) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION

    normalized_pages = tuple(normalize_ocr_page(page) for page in pages)
    document = Document()
    prepared_pages = []
    for page in normalized_pages:
        body_blocks, footer_blocks = split_footer_blocks(page)
        prepared_pages.append(({**page, "blocks": body_blocks}, footer_blocks, page))
    page_scales = [calculate_page_text_scale(body_page) for body_page, _, _ in prepared_pages]
    document_scale = min(page_scales, default=1.0) * 0.92
    document_font_height = typical_document_font_height([body_page for body_page, _, _ in prepared_pages])
    footer_payloads: list[tuple[list[dict[str, Any]], dict[str, Any]]] = []
    for page_offset, (body_page, footer_blocks, page) in enumerate(prepared_pages):
        if page_offset == 0:
            section = document.sections[0]
        else:
            section = document.add_section(WD_SECTION.NEW_PAGE)
            minimize_section_break_paragraph(document.paragraphs[-1])
        configure_ocr_section(section, page)
        add_ocr_page(document, body_page, document_scale, document_font_height)
        footer_payloads.append((footer_blocks, page))
    for section, (footer_blocks, page) in zip(document.sections, footer_payloads):
        set_section_footer(section, footer_blocks, page)
    document.save(output)


def minimize_section_break_paragraph(paragraph: Any) -> None:
    from docx.shared import Pt

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.add_run().font.size = Pt(1)


def configure_ocr_section(section: Any, page: dict[str, Any]) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt

    width = float(page["width"])
    height = float(page["height"])
    section.orientation = WD_ORIENT.LANDSCAPE if width > height else WD_ORIENT.PORTRAIT
    section.page_width = Pt(width)
    section.page_height = Pt(height)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(36)
    section.right_margin = Pt(36)
    section.footer_distance = Pt(18)


def split_footer_blocks(page: dict[str, Any]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    blocks = list(page.get("blocks", []))
    if not blocks:
        return [], []
    page_height = float(page["height"])
    page_width = float(page["width"])
    heights = sorted(float(block.get("height", 1)) for block in blocks)
    typical_height = heights[len(heights) // 2]
    candidates = [
        block for block in blocks
        if float(block.get("y", 0)) + float(block.get("height", 1)) >= page_height * 0.9
        and float(block.get("width", 1)) <= page_width * 0.35
    ]
    if not candidates:
        return blocks, []
    candidate_ids = {id(block) for block in candidates}
    body_bottom = max(
        (float(block.get("y", 0)) + float(block.get("height", 1)) for block in blocks if id(block) not in candidate_ids),
        default=0.0,
    )
    footer_top = min(float(block.get("y", 0)) for block in candidates)
    strong_footer = all(
        float(block.get("y", 0)) >= page_height * 0.92
        and float(block.get("width", 1)) <= page_width * 0.12
        for block in candidates
    )
    if not strong_footer and footer_top - body_bottom < typical_height * 2.5:
        return blocks, []
    return [block for block in blocks if id(block) not in candidate_ids], candidates


def set_section_footer(section: Any, blocks: list[dict[str, Any]], page: dict[str, Any]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.clear()
    if not blocks:
        return
    lines = group_ocr_lines(blocks)
    paragraph.text = "\t".join(str(line["text"]).strip() for line in lines)
    center = sum(float(line["x"]) + float(line["width"]) / 2 for line in lines) / len(lines)
    page_width = float(page["width"])
    if abs(center - page_width / 2) <= page_width * 0.12:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif center > page_width * 0.68:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "Arial"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
        run.font.size = Pt(max(7.0, min(12.0, max(float(line["height"]) for line in lines) * 1.1)))


def normalize_ocr_page(page: dict[str, Any]) -> dict[str, Any]:
    source_width = max(1.0, float(page.get("width", 612)))
    source_height = max(1.0, float(page.get("height", 792)))
    longest_edge = max(source_width, source_height)
    scale = min(1.0, 842.0 / longest_edge)
    normalized = dict(page)
    normalized["width"] = source_width * scale
    normalized["height"] = source_height * scale
    normalized["blocks"] = [
        {
            **block,
            "x": float(block.get("x", 0)) * scale,
            "y": float(block.get("y", 0)) * scale,
            "width": float(block.get("width", 1)) * scale,
            "height": float(block.get("height", 1)) * scale,
        }
        for block in page.get("blocks", [])
    ]
    return normalized


def calculate_page_text_scale(page: dict[str, Any]) -> float:
    blocks = select_body_blocks(page.get("blocks", []))
    lines = group_ocr_lines(blocks)
    if not lines:
        return 1.0
    content_width = max(72.0, float(page["width"]) - 72.0)
    content_height = max(72.0, float(page["height"]) - 72.0)
    left_edge, right_edge = dominant_body_bounds(lines)
    top_edge = min(line["y"] for line in lines)
    bottom_edge = max(line["y"] + line["height"] for line in lines)
    horizontal_scale = content_width / max(1.0, right_edge - left_edge)
    vertical_scale = content_height / max(1.0, bottom_edge - top_edge) * 0.85
    return min(2.5, horizontal_scale, vertical_scale)


def typical_document_font_height(pages: list[dict[str, Any]]) -> float:
    heights = [
        float(line["height"])
        for page in pages
        for line in group_ocr_lines(select_body_blocks(page.get("blocks", [])))
    ]
    if not heights:
        return 12.0
    ordered = sorted(heights)
    return ordered[len(ordered) // 2]


def add_ocr_page(document: Any, page: dict[str, Any], text_scale: float, typical_font_height: float) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    blocks = select_body_blocks(page.get("blocks", []))
    lines = group_ocr_lines(blocks)
    if not lines:
        return
    left_edge, right_edge = dominant_body_bounds(lines)
    paragraphs = group_ocr_paragraphs(lines, left_edge, right_edge)
    previous_bottom = min(line["y"] for line in lines)
    for paragraph_lines in paragraphs:
        first = paragraph_lines[0]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(min(18.0, max(0.0, first["y"] - previous_bottom) * text_scale))
        paragraph.paragraph_format.first_line_indent = Pt(max(0.0, first["x"] - left_edge) * text_scale)
        text = join_wrapped_lines(paragraph_lines, left_edge, right_edge)
        run = paragraph.add_run(text)
        run.font.name = "Arial"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "PingFang SC")
        run.font.bold = False
        recognized_height = max(float(line["height"]) for line in paragraph_lines)
        if typical_font_height * 0.8 <= recognized_height <= typical_font_height * 1.25:
            recognized_height = typical_font_height
        font_size = max(7.0, min(48.0, recognized_height * 1.25 * text_scale))
        run.font.size = Pt(font_size)
        alignment = paragraph_alignment(paragraph_lines, left_edge, right_edge)
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = Pt(line_pitch(paragraph_lines) * text_scale)
        previous_bottom = max(line["y"] + line["height"] for line in paragraph_lines)


def group_ocr_lines(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    lines: list[list[dict[str, Any]]] = []
    ordered = sorted(blocks, key=lambda block: (float(block.get("y", 0)), float(block.get("x", 0))))
    for block in ordered:
        center = float(block.get("y", 0)) + float(block.get("height", 1)) / 2
        matching = next((line for line in lines if abs(center - line_center(line)) <= line_height(line) * 0.45), None)
        if matching is None:
            lines.append([block])
        else:
            matching.append(block)
    return [merge_line_blocks(line) for line in lines]


def select_body_blocks(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    blocks = remove_overlapping_duplicates(blocks)
    if len(blocks) < 5:
        return blocks
    ordered_left = sorted(float(block.get("x", 0)) for block in blocks)
    ordered_right = sorted(float(block.get("x", 0)) + float(block.get("width", 1)) for block in blocks)
    ordered_height = sorted(float(block.get("height", 1)) for block in blocks)
    middle = len(blocks) // 2
    body_left = ordered_left[middle]
    body_right = ordered_right[middle]
    separation = max(12.0, ordered_height[middle] * 2.0)
    selected = []
    for block in blocks:
        left = float(block.get("x", 0))
        right = left + float(block.get("width", 1))
        gap = body_left - right if right < body_left else left - body_right if left > body_right else 0.0
        if gap > separation or detached_from_same_row_body(block, blocks, body_left, separation):
            continue
        selected.append(block)
    return selected or blocks


def remove_overlapping_duplicates(blocks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    selected = []
    for block in blocks:
        text = str(block.get("text", "")).strip()
        left = float(block.get("x", 0))
        right = left + float(block.get("width", 1))
        center = float(block.get("y", 0)) + float(block.get("height", 1)) / 2
        duplicate = False
        for peer in blocks:
            if peer is block or float(peer.get("width", 1)) <= float(block.get("width", 1)) * 1.8:
                continue
            peer_text = str(peer.get("text", "")).strip()
            peer_center = float(peer.get("y", 0)) + float(peer.get("height", 1)) / 2
            block_top = float(block.get("y", 0))
            block_bottom = block_top + float(block.get("height", 1))
            peer_top = float(peer.get("y", 0))
            peer_bottom = peer_top + float(peer.get("height", 1))
            vertical_overlap = max(0.0, min(block_bottom, peer_bottom) - max(block_top, peer_top))
            if vertical_overlap == 0 and abs(center - peer_center) > max(float(block.get("height", 1)), float(peer.get("height", 1))):
                continue
            peer_left = float(peer.get("x", 0))
            peer_right = peer_left + float(peer.get("width", 1))
            overlap = max(0.0, min(right, peer_right) - max(left, peer_left))
            if overlap >= float(block.get("width", 1)) * 0.6 and text and text in peer_text:
                duplicate = True
                break
        if not duplicate:
            selected.append(block)
    return selected


def detached_from_same_row_body(
    block: dict[str, Any],
    blocks: list[dict[str, Any]],
    body_left: float,
    separation: float,
) -> bool:
    block_left = float(block.get("x", 0))
    if abs(block_left - body_left) <= separation:
        return False
    block_center = float(block.get("y", 0)) + float(block.get("height", 1)) / 2
    for peer in blocks:
        if peer is block or abs(float(peer.get("x", 0)) - body_left) > separation:
            continue
        peer_center = float(peer.get("y", 0)) + float(peer.get("height", 1)) / 2
        if abs(block_center - peer_center) > max(float(block.get("height", 1)), float(peer.get("height", 1))) * 0.45:
            continue
        peer_right = float(peer.get("x", 0)) + float(peer.get("width", 1))
        if block_left - peer_right > separation:
            return True
    return False


def dominant_body_bounds(lines: list[dict[str, Any]]) -> tuple[float, float]:
    starts = sorted(float(line["x"]) for line in lines)
    widths = sorted(float(line["width"]) for line in lines)
    heights = sorted(float(line["height"]) for line in lines)
    middle = len(lines) // 2
    dominant_left = starts[middle]
    tolerance = max(10.0, heights[middle] * 1.5)
    typical_width = max(1.0, widths[middle])
    body_lines = [
        line for line in lines
        if abs(float(line["x"]) - dominant_left) <= tolerance
        or float(line["width"]) >= typical_width * 0.85
    ]
    left = min(float(line["x"]) for line in body_lines)
    right = max(float(line["x"]) + float(line["width"]) for line in body_lines)
    return left, right


def merge_line_blocks(blocks: list[dict[str, Any]]) -> dict[str, Any]:
    ordered = sorted(blocks, key=lambda block: float(block.get("x", 0)))
    text = str(ordered[0].get("text", ""))
    right = float(ordered[0].get("x", 0)) + float(ordered[0].get("width", 1))
    for block in ordered[1:]:
        block_x = float(block.get("x", 0))
        average_character = max(2.0, float(block.get("width", 1)) / max(1, len(str(block.get("text", "")))))
        separator = "\t" if block_x - right > average_character * 2.5 else " "
        text += separator + str(block.get("text", ""))
        right = max(right, block_x + float(block.get("width", 1)))
    x = min(float(block.get("x", 0)) for block in ordered)
    y = min(float(block.get("y", 0)) for block in ordered)
    return {"text": text, "x": x, "y": y, "width": right - x, "height": line_height(ordered)}


def group_ocr_paragraphs(
    lines: list[dict[str, Any]],
    content_left: float | None = None,
    content_right: float | None = None,
) -> list[list[dict[str, Any]]]:
    paragraphs: list[list[dict[str, Any]]] = []
    left_edge = content_left if content_left is not None else min(float(line["x"]) for line in lines)
    right_edge = content_right if content_right is not None else max(float(line["x"]) + float(line["width"]) for line in lines)
    typical_height = sorted(float(line["height"]) for line in lines)[len(lines) // 2]
    content_span = max(1.0, right_edge - left_edge)
    for line in lines:
        if not paragraphs:
            paragraphs.append([line])
            continue
        previous = paragraphs[-1][-1]
        gap = float(line["y"]) - (float(previous["y"]) + float(previous["height"]))
        starts_indented_paragraph = float(line["x"]) - left_edge >= typical_height * 0.75
        has_paragraph_gap = gap > max(float(line["height"]), float(previous["height"])) * 1.5
        previous_alignment = line_alignment(previous, left_edge, right_edge)
        current_alignment = line_alignment(line, left_edge, right_edge)
        alignment_changed = previous_alignment != current_alignment and (previous_alignment != "left" or current_alignment != "left")
        starts_at_body_left = abs(float(line["x"]) - left_edge) <= typical_height * 0.75
        previous_right_gap = right_edge - (float(previous["x"]) + float(previous["width"]))
        geometric_boundary = starts_at_body_left and previous_right_gap > content_span * 0.1
        if not starts_indented_paragraph and not has_paragraph_gap and not alignment_changed and not geometric_boundary:
            paragraphs[-1].append(line)
        else:
            paragraphs.append([line])
    return paragraphs


def join_wrapped_lines(lines: list[dict[str, Any]], content_left: float, content_right: float) -> str:
    content_span = max(1.0, content_right - content_left)
    text = ""
    for index, line in enumerate(lines):
        addition = str(line["text"]).strip()
        if text and addition:
            previous = lines[index - 1]
            if line_requires_break(previous, content_left, content_right, content_span):
                text += "\n"
            elif text[-1].isascii() and addition[0].isascii():
                text += " "
        text += addition
    return text


def line_requires_break(line: dict[str, Any], content_left: float, content_right: float, span: float) -> bool:
    right_gap = content_right - (float(line["x"]) + float(line["width"]))
    return right_gap > span * 0.18 or line_alignment(line, content_left, content_right) != "left"


def paragraph_alignment(lines: list[dict[str, Any]], content_left: float, content_right: float) -> str:
    alignments = [line_alignment(line, content_left, content_right) for line in lines]
    for candidate in ("center", "right"):
        if alignments.count(candidate) > len(alignments) / 2:
            return candidate
    return "left"


def line_alignment(line: dict[str, Any], content_left: float, content_right: float) -> str:
    span = max(1.0, content_right - content_left)
    width = float(line["width"])
    left_gap = float(line["x"]) - content_left
    right_gap = content_right - (float(line["x"]) + width)
    tolerance = max(8.0, span * 0.08)
    if width < span * 0.8 and abs(left_gap - right_gap) <= tolerance:
        return "center"
    if width < span * 0.8 and left_gap > max(tolerance, right_gap * 1.5):
        return "right"
    return "left"


def line_pitch(lines: list[dict[str, Any]]) -> float:
    if len(lines) < 2:
        return max(12.0, float(lines[0]["height"]) * 1.35)
    pitches = [
        float(current["y"]) - float(previous["y"])
        for previous, current in zip(lines, lines[1:])
        if float(current["y"]) > float(previous["y"])
    ]
    if not pitches:
        return max(12.0, float(lines[0]["height"]) * 1.35)
    return max(12.0, sorted(pitches)[len(pitches) // 2])


def line_center(blocks: list[dict[str, Any]]) -> float:
    return sum(float(block.get("y", 0)) + float(block.get("height", 1)) / 2 for block in blocks) / len(blocks)


def line_height(blocks: list[dict[str, Any]]) -> float:
    return max(1.0, max(float(block.get("height", 1)) for block in blocks))


def main() -> int:
    try:
        convert(read_request())
        return 0
    except Exception as error:
        emit("failed", code=error_code(error))
        traceback.print_exc(file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
