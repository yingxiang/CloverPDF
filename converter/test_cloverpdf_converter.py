from __future__ import annotations

import io
import os
import sys
import tempfile
import unittest
from types import SimpleNamespace
from unittest.mock import patch

from converter import cloverpdf_converter


class ConverterWorkerTests(unittest.TestCase):
    def test_multi_page_ocr_uses_page_break_before_without_empty_paragraph(self) -> None:
        from docx import Document

        pages = (
            {"page": 0, "width": 612, "height": 792, "blocks": [{"text": "Page one", "x": 10, "y": 10, "width": 100, "height": 12}]},
            {"page": 1, "width": 612, "height": 792, "blocks": [{"text": "Page two", "x": 10, "y": 10, "width": 100, "height": 12}]},
        )
        with tempfile.TemporaryDirectory() as directory:
            output = os.path.join(directory, "output.docx")
            cloverpdf_converter.build_ocr_document(pages, output)
            document = Document(output)

        self.assertEqual([paragraph.text for paragraph in document.paragraphs], ["Page one", "Page two"])
        self.assertTrue(document.paragraphs[1].paragraph_format.page_break_before)

    def test_geometry_breaks_lines_that_return_to_body_left(self) -> None:
        lines = [
            {"text": "long wrapped line", "x": 10, "y": 0, "width": 79, "height": 10},
            {"text": "continuation", "x": 10, "y": 14, "width": 78, "height": 10},
            {"text": "short line", "x": 10, "y": 28, "width": 45, "height": 10},
            {"text": "new line", "x": 10, "y": 42, "width": 75, "height": 10},
        ]

        paragraphs = cloverpdf_converter.group_ocr_paragraphs(lines, 10, 90)

        self.assertEqual([[line["text"] for line in paragraph] for paragraph in paragraphs], [
            ["long wrapped line", "continuation", "short line"],
            ["new line"],
        ])

    def test_overlapping_stamp_text_is_removed_when_body_already_contains_it(self) -> None:
        blocks = [
            {"text": "上海进馨网路科技有限公司", "x": 100, "y": 100, "width": 200, "height": 20},
            {"text": "馨", "x": 150, "y": 105, "width": 30, "height": 18},
        ]

        selected = cloverpdf_converter.remove_overlapping_duplicates(blocks)

        self.assertEqual([block["text"] for block in selected], ["上海进馨网路科技有限公司"])

    def test_detached_block_on_body_row_is_removed_before_line_merge(self) -> None:
        blocks = [
            {"text": f"body-{index}", "x": 20, "y": index * 15, "width": 70, "height": 10}
            for index in range(6)
        ]
        blocks.append({"text": "stamp", "x": 130, "y": 30, "width": 80, "height": 10})

        selected = cloverpdf_converter.select_body_blocks(blocks)
        lines = cloverpdf_converter.group_ocr_lines(selected)
        left, right = cloverpdf_converter.dominant_body_bounds(lines)

        self.assertEqual((left, right), (20, 90))
        self.assertNotIn("stamp", " ".join(line["text"] for line in lines))

    def test_detached_edge_block_does_not_expand_body_bounds(self) -> None:
        blocks = [
            {"text": f"body-{index}", "x": 20, "y": index * 15, "width": 70, "height": 10}
            for index in range(6)
        ]
        blocks.append({"text": "stamp", "x": 130, "y": 40, "width": 65, "height": 10})

        selected = cloverpdf_converter.select_body_blocks(blocks)

        self.assertEqual(len(selected), 6)
        self.assertNotIn("stamp", [block["text"] for block in selected])

    def test_centered_heading_is_separate_from_left_body(self) -> None:
        lines = [
            {"text": "Title", "x": 40, "y": 0, "width": 20, "height": 10},
            {"text": "Body", "x": 10, "y": 14, "width": 80, "height": 10},
        ]

        paragraphs = cloverpdf_converter.group_ocr_paragraphs(lines, 10, 90)

        self.assertEqual([len(paragraph) for paragraph in paragraphs], [1, 1])

    def test_ocr_blocks_use_spaces_tabs_and_paragraph_breaks(self) -> None:
        blocks = [
            {"text": "A", "x": 10, "y": 10, "width": 10, "height": 10},
            {"text": "B", "x": 24, "y": 10, "width": 10, "height": 10},
            {"text": "C", "x": 80, "y": 10, "width": 10, "height": 10},
            {"text": "next", "x": 10, "y": 24, "width": 30, "height": 10},
            {"text": "paragraph", "x": 20, "y": 65, "width": 50, "height": 10},
        ]

        lines = cloverpdf_converter.group_ocr_lines(blocks)
        paragraphs = cloverpdf_converter.group_ocr_paragraphs(lines)

        self.assertEqual(lines[0]["text"], "A B\tC")
        self.assertEqual([len(paragraph) for paragraph in paragraphs], [2, 1])
        self.assertEqual(cloverpdf_converter.join_wrapped_lines(paragraphs[0], 10, 90), "A B\tC next")

    def test_alignment_and_manual_break_use_normalized_content_bounds(self) -> None:
        centered = {"text": "Title", "x": 40, "y": 0, "width": 20, "height": 10}
        right = {"text": "Date", "x": 70, "y": 0, "width": 20, "height": 10}
        full = {"text": "Body", "x": 10, "y": 0, "width": 80, "height": 10}

        self.assertEqual(cloverpdf_converter.line_alignment(centered, 10, 90), "center")
        self.assertEqual(cloverpdf_converter.line_alignment(right, 10, 90), "right")
        self.assertEqual(cloverpdf_converter.line_alignment(full, 10, 90), "left")
        self.assertTrue(cloverpdf_converter.line_requires_break(centered, 10, 90, 80))
        self.assertFalse(cloverpdf_converter.line_requires_break(full, 10, 90, 80))

    def test_normalize_ocr_page_preserves_portrait_aspect_ratio(self) -> None:
        page = {
            "width": 3072,
            "height": 4096,
            "blocks": [{"text": "Title", "x": 307.2, "y": 409.6, "width": 1536, "height": 200}],
        }

        normalized = cloverpdf_converter.normalize_ocr_page(page)

        self.assertAlmostEqual(normalized["width"], 631.5)
        self.assertAlmostEqual(normalized["height"], 842.0)
        self.assertAlmostEqual(normalized["width"] / normalized["height"], 0.75)
        self.assertAlmostEqual(normalized["blocks"][0]["x"], 63.15)
        self.assertAlmostEqual(normalized["blocks"][0]["width"], 315.75)

    def test_read_request_accepts_non_contiguous_pages(self) -> None:
        payload = '{"input":"in.pdf","output":"out.docx","password":null,"pages":[0,2]}\n'
        with patch.object(sys, "stdin", io.StringIO(payload)):
            request = cloverpdf_converter.read_request()

        self.assertEqual(request.pages, (0, 2))
        self.assertIsNone(request.ocr_pages)

    def test_convert_passes_selected_pages_to_pdf2docx(self) -> None:
        calls: list[dict[str, object]] = []

        class FakeConverter:
            def __init__(self, input_path: str, password: str) -> None:
                self.input_path = input_path
                self.password = password

            def convert(self, output_path: str, **kwargs: object) -> None:
                calls.append(kwargs)
                with open(output_path, "wb") as output:
                    output.write(b"docx")

            def close(self) -> None:
                pass

        with tempfile.TemporaryDirectory() as directory:
            input_path = os.path.join(directory, "input.pdf")
            output_path = os.path.join(directory, "output.docx")
            with open(input_path, "wb") as source:
                source.write(b"pdf")
            request = cloverpdf_converter.Request(
                input=input_path,
                output=output_path,
                password=None,
                pages=(0, 2),
                ocr_pages=None,
            )
            fake_module = SimpleNamespace(Converter=FakeConverter)
            with patch.dict(sys.modules, {"pdf2docx": fake_module}):
                cloverpdf_converter.convert(request)

        self.assertEqual(calls, [{"multi_processing": False, "pages": [0, 2]}])

    def test_convert_ocr_writes_each_recognized_page(self) -> None:
        saved: list[tuple[tuple[dict[str, object], ...], str]] = []

        def fake_build(pages: tuple[dict[str, object], ...], path: str) -> None:
            saved.append((pages, path))
            with open(path, "wb") as output:
                output.write(b"docx")

        with tempfile.TemporaryDirectory() as directory:
            input_path = os.path.join(directory, "input.pdf")
            output_path = os.path.join(directory, "output.docx")
            with open(input_path, "wb") as source:
                source.write(b"pdf")
            request = cloverpdf_converter.Request(
                input=input_path,
                output=output_path,
                password=None,
                pages=(0, 2),
                ocr_pages=(
                    {"page": 0, "width": 612, "height": 792, "blocks": [{"text": "Title"}]},
                    {"page": 2, "width": 612, "height": 792, "blocks": [{"text": "Body"}]},
                ),
            )
            with patch.object(cloverpdf_converter, "build_ocr_document", fake_build):
                cloverpdf_converter.convert(request)

        self.assertEqual(saved[0][1], output_path)
        self.assertEqual([page["page"] for page in saved[0][0]], [0, 2])


if __name__ == "__main__":
    unittest.main()
